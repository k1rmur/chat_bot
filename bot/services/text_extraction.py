import logging
import os

import docx
import textract
import langchain_core.documents
from aiogram import Bot
from aiogram.types import Document

logger = logging.getLogger(__name__)


async def extract_text_from_document(document: Document, bot: Bot):
    """
    Функция, переводящая объект Document aiogram (который содержит docx файл) в Document Langchain
    """

    file_info = await bot.get_file(document.file_id)
    file_path = file_info.file_path

    downloaded_file = await bot.download_file(file_path)

    temp_file_name = f"temp_{document.file_id}.docx"

    with open(temp_file_name, "wb") as f:
        f.write(downloaded_file.getvalue())

    full_text = textract.process(temp_file_name)

    os.remove(temp_file_name)

    langchain_document = langchain_core.documents.Document(
        page_content=full_text
    )

    return langchain_document


async def extract_news_from_document(document: Document, bot: Bot):
    """
    Функция, переводящая объект Document aiogram (который содержит docx файл) в Documentы Langchain
    """

    file_info = await bot.get_file(document.file_id)
    file_path = file_info.file_path

    downloaded_file = await bot.download_file(file_path)

    temp_file_name = f"temp_{document.file_id}.docx"

    with open(temp_file_name, "wb") as f:
        f.write(downloaded_file.getvalue())

    doc = docx.Document(temp_file_name)

    headings = [
        "факты и события. в россии",
        "в мире",
        "конференции",
        "периодические издания",
    ]
    paragraphs = doc.paragraphs

    indices = []
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if paragraph.style.name.startswith("Heading"):
            if paragraph.text.lower() in headings:
                indices.append(i)

    facts = "\n".join([par.text for par in paragraphs[indices[0] : indices[1]]])
    world = "\n".join([par.text for par in paragraphs[indices[1] : indices[2]]])
    conferences = "\n".join([par.text for par in paragraphs[indices[2] : indices[3]]])

    os.remove(temp_file_name)

    return tuple(
        langchain_core.documents.Document(page_content="\n".join(text))
        for text in (facts, world, conferences)
    )
