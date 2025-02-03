import asyncio
import logging
import os

import torch
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from dotenv import find_dotenv, load_dotenv
from langchain.chains.llm import LLMChain
from langchain_community.chat_models import GigaChat
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import CharacterTextSplitter

from .prompt_templates import (
    CONTENT_PROMPT,
    EXTRACT_MAIN_QUESTIONS_PROMPT,
    EXTRACT_SHORT_RESUME_PROMPT,
    ORDER_PROMPT,
    REFINE_CONTENT_PROMPT,
    REFINE_ORDERS_PROMPT,
)

logger = logging.getLogger(__name__)
CONTEXT_LENGTH = 8192 * 4

load_dotenv(find_dotenv())


async def process_chunk(order_chain, content_chain, chunks):
    orders = []
    contents = []
    for chunk in chunks:
        orders.append(await order_chain.ainvoke({"chunk": chunk}))
        contents.append(await content_chain.ainvoke({"chunk": chunk}))
    return orders, contents


ABS_PATH = os.path.dirname(os.path.abspath(__file__))

device = "cuda" if torch.cuda.is_available() else "cpu"
print("DEVICE:", device)

llm_nonrag = GigaChat(verify_ssl_certs=False, credentials=os.getenv("CREDENTIALS_NONRAG"), scope="GIGACHAT_API_CORP", model="GigaChat")

def length_function(text) -> int:
    """Get number of tokens for input contents."""
    return llm_nonrag.get_num_tokens(text)


REFINE_ORDERS_PROMPT_LENGTH = length_function(REFINE_ORDERS_PROMPT)
REFINE_CONTENT_PROMPT_LENGTH = length_function(REFINE_CONTENT_PROMPT)
EXTRACT_MAIN_QUESTIONS_PROMPT_LENGTH = length_function(EXTRACT_MAIN_QUESTIONS_PROMPT)
EXTRACT_SHORT_RESUME_PROMPT_LENGTH = length_function(EXTRACT_SHORT_RESUME_PROMPT)
CONTENT_PROMPT_LENGTH = length_function(CONTENT_PROMPT)
ORDER_PROMPT_LENGTH = length_function(ORDER_PROMPT)


refine_orders_prompt = PromptTemplate.from_template(REFINE_ORDERS_PROMPT)

refine_content_prompt = PromptTemplate.from_template(REFINE_CONTENT_PROMPT)

extract_main_questions_prompt = PromptTemplate.from_template(
    EXTRACT_MAIN_QUESTIONS_PROMPT
)

extract_short_resume_prompt = PromptTemplate.from_template(EXTRACT_SHORT_RESUME_PROMPT)


content_template = PromptTemplate(input_variables=["chunk"], template=CONTENT_PROMPT)
order_template = PromptTemplate(input_variables=["chunk"], template=ORDER_PROMPT)


async def get_summary(file_id, text, message):
    prompt_length = max(CONTENT_PROMPT_LENGTH, ORDER_PROMPT_LENGTH)

    text_splitter = CharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=CONTEXT_LENGTH - prompt_length - 500, chunk_overlap=100
    )

    chunks = text_splitter.split_text(text)

    logger.info(f'Чанки: {" ".join(chunks)}')

    order_chain = order_template | llm_nonrag
    content_chain = content_template | llm_nonrag
    orders, contents = await process_chunk(order_chain, content_chain, chunks)

    raw_orders = []
    raw_content = []
    for ord in orders:
        if ord:
            ord = ord.content
            raw_orders.append(ord)
    for cont in contents:
        if cont:
            cont = cont.content
            raw_content.append(cont)

    raw_orders = "\n".join(raw_orders)
    raw_content = "\n".join(raw_content)

    logger.info(f"raw_orders: {raw_orders}\n{len(raw_orders)}")
    logger.info(f"raw_content: {raw_content}\n{len(raw_content)}")

    recursion_content = 0
    while True:
        content_chunks = text_splitter.split_text(raw_content)
        tasks = [content_chain.ainvoke({"chunk": chunk}) for chunk in content_chunks]
        results = await asyncio.gather(*tasks)
        raw_content = "\n".join([cont.content for cont in results])
        logger.info(f"refined_content: {raw_content}\n{len(raw_content)}")
        recursion_content += 1
        if (
            length_function(raw_content)
            < CONTEXT_LENGTH - REFINE_CONTENT_PROMPT_LENGTH - 500
            or recursion_content > 10
        ):
            break

    recursion_orders = 0
    while True:
        order_chunks = text_splitter.split_text(raw_orders)
        tasks = [order_chain.ainvoke({"chunk": chunk}) for chunk in order_chunks]
        results = await asyncio.gather(*tasks)
        raw_orders = "\n".join([ord.content for ord in results])
        logger.info(f"refined_orders: {raw_orders}\n{len(raw_orders)}")
        recursion_orders += 1
        if (
            length_function(raw_orders)
            < CONTEXT_LENGTH - REFINE_CONTENT_PROMPT_LENGTH - 500
            or recursion_orders > 10
        ):
            break

    content = await LLMChain(llm=llm_nonrag, prompt=refine_content_prompt).ainvoke(raw_content)
    content = content["text"]
    orders = await LLMChain(llm=llm_nonrag, prompt=refine_orders_prompt).ainvoke(raw_orders)
    orders = orders["text"]

    logger.info(f"orders: {orders}")
    logger.info(f"content: {content}")

    main_question = await LLMChain(
        llm=llm_nonrag, prompt=extract_main_questions_prompt
    ).ainvoke(
        f"Темы разговора и принятые решения:{content}\nПоручения сотрудникам:{orders}"
    )
    resume = await LLMChain(llm=llm_nonrag, prompt=extract_short_resume_prompt).ainvoke(
        f"Темы разговора и принятые решения:{content}\nПоручения сотрудникам:{orders}"
    )

    doc = Document()
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = "Times New Roman"

    heading = doc.add_paragraph()
    heading.add_run("Основной вопрос", style="CommentsStyle").bold = True
    heading = doc.add_paragraph()
    heading.add_run(f'{main_question["text"]}', style="CommentsStyle")
    doc.add_paragraph("")

    heading = doc.add_paragraph()
    heading.add_run("Краткое резюме", style="CommentsStyle").bold = True
    heading = doc.add_paragraph()
    heading.add_run(f'{resume["text"]}', style="CommentsStyle")
    doc.add_paragraph("")

    heading = doc.add_paragraph()
    heading.add_run("Поручения", style="CommentsStyle").bold = True
    heading = doc.add_paragraph()
    heading.add_run(f"{orders}", style="CommentsStyle")
    doc.add_paragraph("")

    heading = doc.add_paragraph()
    heading.add_run("Содержание встречи", style="CommentsStyle").bold = True
    heading = doc.add_paragraph()
    heading.add_run(content, style="CommentsStyle")

    doc.save(f"./tmp/{file_id}.docx")

    return f"./tmp/{file_id}.docx", "Протокол.docx"
