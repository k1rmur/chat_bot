import asyncio
import functools
import glob
import os

import moviepy.editor as mp
import whisperx
from docx import Document
from dotenv import find_dotenv, load_dotenv
from pydub import AudioSegment
from salute_speech.speech_recognition import SaluteSpeechClient
from salute_speech.speech_recognition import SpeechRecognitionConfig


load_dotenv(find_dotenv())


device = "cpu"
batch_size = 16

compute_type = "int8"
model_dir = "/app/bot/services/models/"
model_name = "medium"

model = whisperx.load_model(
    model_name,
    device=device,
    compute_type=compute_type,
    download_root=model_dir,
    language="ru",
)


def change_speaker_name(string: str | None):
    if string is None:
        return "Собеседник"
    else:
        spk, num = string.split("_")
        num = int(num) + 1
    return f"Собеседник {num}"


HF_TOKEN = os.getenv("HF_TOKEN")


def clear_temp(file_id):
    files = glob.glob("./tmp/*")
    for f in files:
        if file_id in f:
            print(f"Deleting {f}...")
            os.remove(f)
    files = glob.glob("app/bot/tmp/*")
    for f in files:
        if file_id in f:
            print(f"Deleting {f}...")
            os.remove(f)


def is_video(extension):
    extention_list = ["webm", "mp4", "mkv", "flv", "avi", "mov", "wmv", "m4v"]

    return extension.lower() in extention_list


def is_audio(extension):
    extension_list = [
        "mp3",
        "oga",
        "ogg",
        "mogg",
        "3gp",
        "aac",
        "aa",
        "aax",
        "m4a",
        "mpc",
        "wav",
    ]

    return extension.lower() in extension_list


def convert(video_path, audio_path):
    clip = mp.VideoFileClip(video_path)
    clip.audio.write_audiofile(audio_path, ffmpeg_params=["-ac", "1"])


lock = asyncio.Lock()


async def salute_recognize(file_id: str, extension: str):

    audio_path = f"/app/bot/tmp/{file_id}.{extension}"
    client = SaluteSpeechClient(client_credentials=os.getenv("SBER_SPEECH_API_KEY"))

    if extension not in ["mp3", "wav"]:
        song = AudioSegment.from_ogg(f"/app/bot/tmp/{file_id}.{extension}")
        song.export(f"/app/bot/tmp/{file_id}.mp3", format="mp3")

        audio_path = f"/app/bot/tmp/{file_id}.mp3"

    text_file = f"/app/bot/tmp/{file_id}.txt"

    with open(audio_path, "rb") as audio_file:
        result = await client.audio.transcriptions.create(
            file=audio_file,
            language="ru-RU",
        )

    doc_transcription = Document()

    doc_transcription.add_paragraph(f"{result.text}")

    doc_transcription.save(f"/app/bot/tmp/{file_id}.docx")

    return f"/app/bot/tmp/{file_id}.docx", "Транскрипция.docx", result.text


async def recognize(file_id: str, extension: str) -> None:
    loop = asyncio.get_event_loop()
    audiofile = f"/app/bot/tmp/{file_id}.{extension}"
    audio = whisperx.load_audio(audiofile)

    transcription_result = await loop.run_in_executor(
        None,
        functools.partial(
            model.transcribe, audiofile, language="ru", batch_size=batch_size
        ),
    )
    model_a, metadata = whisperx.load_align_model(language_code="ru", device=device)
    result = await loop.run_in_executor(
        None,
        functools.partial(
            whisperx.align,
            transcription_result["segments"],
            model_a,
            metadata,
            audio,
            device,
            return_char_alignments=False,
        ),
    )
    diarize_model = whisperx.DiarizationPipeline(
        use_auth_token=HF_TOKEN,
        device=device,
        model_name="pyannote/speaker-diarization-3.1",
    )
    diarize_segments = await loop.run_in_executor(
        None, functools.partial(diarize_model, audio)
    )
    result = whisperx.assign_word_speakers(diarize_segments, result)

    text_for_summary = []
    doc_transcription = Document()

    for segment in result["segments"]:
        speaker = change_speaker_name(segment.get("speaker"))
        line_massive = f"[{segment.get('start'):.3f} --> {segment.get('end'):.3f}] {speaker}: {segment.get('text')}"
        line_summary = f"{speaker.split()[-1]}: {segment.get('text')}\n"
        text_for_summary.append(line_summary)
        doc_transcription.add_paragraph(line_massive)

    doc_transcription.save(f"/app/bot/tmp/{file_id}.docx")

    full_transcript = "\n".join(text_for_summary)

    with open(f"/app/bot/tmp/{file_id}.txt", "w") as file:
        file.write(full_transcript)

    return f"/app/bot/tmp/{file_id}.docx", "Транскрипция.docx", full_transcript
